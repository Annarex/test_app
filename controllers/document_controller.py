"""
Контроллер для работы с Word документами (заключения, письма)
Реализует логику из 1С: ЗаключениеКВ
"""
from PyQt5.QtCore import QObject, pyqtSignal
from typing import Dict, Any, Optional, List
from datetime import datetime
from pathlib import Path
import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re
from logger import logger

from models.database import DatabaseManager
from models.base_models import Project


class DocumentController(QObject):
    """Контроллер для формирования документов Word"""
    
    # Сигналы
    document_generated = pyqtSignal(str)  # путь к сгенерированному файлу
    error_occurred = pyqtSignal(str)
    
    def __init__(self, db_manager: DatabaseManager):
        super().__init__()
        self.db_manager = db_manager
        self.templates_dir = Path("templates")
    
    def generate_conclusion(
        self,
        project_id: int,
        revision_id: int,
        protocol_date: datetime,
        protocol_number: str,
        letter_date: Optional[datetime] = None,
        letter_number: Optional[str] = None,
        admin_date: Optional[datetime] = None,
        admin_number: Optional[str] = None,
        output_path: Optional[str] = None
    ) -> Optional[str]:
        """
        Формирование заключения на основе данных формы
        
        Args:
            project_id: ID проекта
            revision_id: ID ревизии формы
            protocol_date: Дата протокола
            protocol_number: Номер протокола
            letter_date: Дата письма (опционально)
            letter_number: Номер письма (опционально)
            admin_date: Дата постановления администрации (опционально)
            admin_number: Номер постановления администрации (опционально)
            output_path: Путь для сохранения файла (опционально)
        
        Returns:
            Путь к сгенерированному файлу или None при ошибке
        """
        try:
            # Загружаем данные ревизии
            revision_data = self.db_manager.load_revision_data(project_id, revision_id)
            if not revision_data:
                self.error_occurred.emit("Данные ревизии не найдены")
                return None
            
            # Загружаем проект
            projects = self.db_manager.load_projects()
            project = next((p for p in projects if p.id == project_id), None)
            if not project:
                self.error_occurred.emit("Проект не найден")
                return None
            
            # Загружаем расширенную информацию о МО
            municipality = None
            if project.municipality_id:
                municipality = self.db_manager.get_municipality_by_id(project.municipality_id)
            
            # Загружаем шаблон
            template_path = self.templates_dir / "шаблон_решения.docx"
            if not template_path.exists():
                self.error_occurred.emit(f"Шаблон не найден: {template_path}")
                return None
            
            doc = docx.Document(str(template_path))
            
            # Извлекаем данные из формы
            form_data = self._extract_form_data(revision_data)
            
            # Заменяем метки в документе
            self._replace_placeholders(
                doc,
                project=project,
                municipality=municipality,
                form_data=form_data,
                protocol_date=protocol_date,
                protocol_number=protocol_number,
                letter_date=letter_date,
                letter_number=letter_number,
                admin_date=admin_date,
                admin_number=admin_number
            )
            
            # Формируем таблицы (Таблица 2, Таблица 3)
            self._insert_table2(doc, form_data, protocol_date)
            self._insert_table3(doc, form_data, protocol_date)
            
            # Сохраняем документ
            if not output_path:
                output_dir = Path("data") / "projects" / str(project_id) / "documents"
                output_dir.mkdir(parents=True, exist_ok=True)
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                output_path = str(output_dir / f"заключение_{timestamp}.docx")
            
            doc.save(output_path)
            self.document_generated.emit(output_path)
            logger.info(f"Заключение сформировано: {output_path}")
            return output_path
            
        except Exception as e:
            error_msg = f"Ошибка формирования заключения: {e}"
            logger.error(error_msg, exc_info=True)
            self.error_occurred.emit(error_msg)
            return None
    
    def generate_letters(
        self,
        project_id: int,
        revision_id: int,
        protocol_date: datetime,
        protocol_number: str,
        output_dir: Optional[str] = None
    ) -> Dict[str, Optional[str]]:
        """
        Формирование писем администрации и совета
        
        Returns:
            Словарь с путями к файлам: {'admin': путь, 'council': путь}
        """
        result = {'admin': None, 'council': None}
        
        try:
            # Загружаем проект
            projects = self.db_manager.load_projects()
            project = next((p for p in projects if p.id == project_id), None)
            if not project:
                self.error_occurred.emit("Проект не найден")
                return result
            
            # Загружаем расширенную информацию о МО
            municipality = None
            if project.municipality_id:
                municipality = self.db_manager.get_municipality_by_id(project.municipality_id)
            
            # Загружаем шаблон
            template_path = self.templates_dir / "шаблон_письма.docx"
            if not template_path.exists():
                self.error_occurred.emit(f"Шаблон не найден: {template_path}")
                return result
            
            # Формируем письмо администрации
            result['admin'] = self._generate_letter_admin(
                template_path, project, municipality, protocol_date, protocol_number, output_dir
            )
            
            # Формируем письмо совета
            result['council'] = self._generate_letter_council(
                template_path, project, municipality, protocol_date, protocol_number, output_dir
            )
            
            return result
            
        except Exception as e:
            error_msg = f"Ошибка формирования писем: {e}"
            logger.error(error_msg, exc_info=True)
            self.error_occurred.emit(error_msg)
            return result
    
    def _extract_form_data(self, revision_data: Dict[str, Any]) -> Dict[str, Any]:
        """Извлечение данных из формы для формирования заключения"""
        доходы_data = revision_data.get('доходы_data', [])
        расходы_data = revision_data.get('расходы_data', [])
        
        # Находим строки "Доходы бюджета - всего" и "Расходы бюджета - всего"
        доходы_всего = None
        расходы_всего = None
        
        for item in доходы_data:
            if 'доходы бюджета - всего' in item.get('наименование_показателя', '').lower():
                доходы_всего = item
                break
        
        for item in расходы_data:
            if 'расходы бюджета - всего' in item.get('наименование_показателя', '').lower():
                расходы_всего = item
                break
        
        # Берем значение из первого столбца бюджета (обычно это "бюджет субъекта Российской Федерации")
        budget_col = 'бюджет субъекта Российской Федерации'
        
        доходы_утвержденный = доходы_всего.get('утвержденный', {}).get(budget_col, 0) if доходы_всего else 0
        доходы_исполненный = доходы_всего.get('исполненный', {}).get(budget_col, 0) if доходы_всего else 0
        расходы_утвержденный = расходы_всего.get('утвержденный', {}).get(budget_col, 0) if расходы_всего else 0
        расходы_исполненный = расходы_всего.get('исполненный', {}).get(budget_col, 0) if расходы_всего else 0

        # Вычисляем дефицит/профицит напрямую из итоговых строк
        дефицит_утвержденный = расходы_утвержденный - доходы_утвержденный
        дефицит_исполненный = расходы_исполненный - доходы_исполненный
        
        # Вычисляем проценты исполнения
        процент_доходов = (доходы_исполненный / доходы_утвержденный * 100) if доходы_утвержденный != 0 else 0
        процент_расходов = (расходы_исполненный / расходы_утвержденный * 100) if расходы_утвержденный != 0 else 0
        
        # Определяем дефицит/профицит
        slojilsya = "профицит" if расходы_исполненный < доходы_исполненный else "дефицит"
        
        return {
            'доходы_всего': доходы_всего,
            'расходы_всего': расходы_всего,
            'доходы_утвержденный': доходы_утвержденный,
            'доходы_исполненный': доходы_исполненный,
            'расходы_утвержденный': расходы_утвержденный,
            'расходы_исполненный': расходы_исполненный,
            'процент_доходов': процент_доходов,
            'процент_расходов': процент_расходов,
            'дефицит_утвержденный': дефицит_утвержденный,
            'дефицит_исполненный': дефицит_исполненный,
            'slojilsya': slojilsya,
            'доходы_data': доходы_data,
            'расходы_data': расходы_data
        }
    
    def _replace_placeholders(
        self,
        doc: docx.Document,
        project: Project,
        municipality: Optional[Any],
        form_data: Dict[str, Any],
        protocol_date: datetime,
        protocol_number: str,
        letter_date: Optional[datetime] = None,
        letter_number: Optional[str] = None,
        admin_date: Optional[datetime] = None,
        admin_number: Optional[str] = None
    ):
        """Замена меток в документе - логика из 1С"""
        # Получаем период из ревизии (TODO: получать из периода ревизии)
        period_name = 'I квартал'  # По умолчанию
        data_sdachi = protocol_date.strftime('%d.%m.') + str(protocol_date.year)
        
        # Определяем совет (городской/муниципальный) по коду родителя МО
        sovet1 = ""
        sovet2 = ""
        if municipality:
            # TODO: проверять код родителя МО из справочника
            # Если код родителя = "00006" -> городской совет
            # Если код родителя = "00007" -> муниципальный совет
            pass
        
        # Вычисляем дополнительные значения из данных формы
        доходы_data = form_data.get('доходы_data', [])
        расходы_data = form_data.get('расходы_data', [])
        
        # Находим максимальный удельный вес для доходов и расходов
        max_udelny_ves_доходы = self._find_max_udelny_ves(доходы_data, form_data['доходы_исполненный'])
        max_udelny_ves_расходы = self._find_max_udelny_ves(расходы_data, form_data['расходы_исполненный'])
        
        # Находим строки с максимальным удельным весом
        max_доход = self._find_item_by_udelny_ves(доходы_data, max_udelny_ves_доходы, form_data['доходы_исполненный'])
        max_расход = self._find_item_by_udelny_ves(расходы_data, max_udelny_ves_расходы, form_data['расходы_исполненный'])
        
        replacements = {
            '<Okrug>': municipality.name if municipality else project.name,
            '<OkrugRP>': getattr(municipality, 'родительный_падеж', municipality.name if municipality else project.name),
            '<Period>': period_name,
            '<Year>': str(protocol_date.year),
            '<DataSdachi>': data_sdachi,
            '<DatePr>': protocol_date.strftime('%d.%m.%Y'),
            '<NomerPr>': protocol_number,
            '<Dohodi2>': self._format_number(form_data['доходы_утвержденный'] / 1000),
            '<Rashodi2>': self._format_number(form_data['расходы_утвержденный'] / 1000),
            '<DohodiISP>': self._format_number(form_data['доходы_исполненный'] / 1000),
            '<RashodiISP>': self._format_number(form_data['расходы_исполненный'] / 1000),
            '<DiffISP>': self._format_number(form_data['дефицит_исполненный'] / 1000),
            '<PRISP1>': self._format_number(form_data['процент_доходов']),
            '<PRISP2>': self._format_number(form_data['процент_расходов']),
            '<Defecit2>': self._format_number(form_data['дефицит_утвержденный'] / 1000),
            '<Defecit3>': self._format_number(abs(form_data['дефицит_утвержденный']) / 1000),
            '<Slojilsya>': form_data['slojilsya'],
            '<Sovet1>': sovet1,
            '<Sovet2>': sovet2,
            '<IzmVs>': self._format_number(form_data.get('izm_vs', 0)),
            '<UvOst>': self._format_number(form_data.get('uv_ost', 0)),
            '<UmOst>': self._format_number(form_data.get('um_ost', 0)),
            '<UvOst1>': self._format_number(form_data.get('uv_ost1', 0)),
            '<UvOst2>': self._format_number(form_data.get('uv_ost2', 0)),
            '<UvOst3>': self._format_number(form_data.get('uv_ost3', 0)),
        }
        
        # Дополнительные метки из кода 1С
        if letter_date:
            replacements['<Kogda1>'] = letter_date.strftime('%d.%m.%Y')
        if letter_number:
            replacements['<Nomer1>'] = letter_number
        if admin_date:
            replacements['<ДатаПостановленияАдм>'] = admin_date.strftime('%d.%m.%Y')
        if admin_number:
            replacements['<НомерПостановленияАдм>'] = admin_number
        
        # Метки для данных из МО (если есть)
        if municipality:
            # TODO: получать из БД
            replacements['<DateZD>'] = getattr(municipality, 'дата_соглашения', '')
            replacements['<DataRYear>'] = getattr(municipality, 'дата_решения', '')
            replacements['<NomerRYear>'] = getattr(municipality, 'номер_решения', '')
            replacements['<SummaDNachalo>'] = self._format_number(getattr(municipality, 'начальная_доходы', 0))
            replacements['<SummaRNachalo>'] = self._format_number(getattr(municipality, 'начальная_расходы', 0))
            начальная_дефицит = getattr(municipality, 'начальная_дефицит', 0)
            if начальная_дефицит == 0:
                replacements['<DefecitNachalo>'] = "0,00"
            else:
                replacements['<DefecitNachalo>'] = self._format_number(начальная_дефицит)
        
        # Находим данные для налоговых и неналоговых доходов (код "10000000000000000")
        # и безвозмездных поступлений (код "20000000000000000")
        nn_dohodi_item = None
        sbpb_item = None
        budget_col = 'бюджет субъекта Российской Федерации'
        
        for item in доходы_data:
            код = item.get('код_классификации', '').replace(' ', '')
            if код == '10000000000000000':
                nn_dohodi_item = item
            elif код == '20000000000000000':
                sbpb_item = item
        
        # Метки для налоговых и неналоговых доходов
        if nn_dohodi_item:
            nn_dohodi_исполнение = nn_dohodi_item.get('исполненный', {}).get(budget_col, 0) / 1000
            nn_dohodi_план = 0
            nn_dohodi_убн = nn_dohodi_item.get('утвержденный', {}).get(budget_col, 0) / 1000
            if nn_dohodi_убн != 0:
                nn_dohodi_план = round(nn_dohodi_исполнение / nn_dohodi_убн * 100, 1)
            nn_dohodi_удельный_вес = 0
            if form_data['доходы_исполненный'] != 0:
                nn_dohodi_удельный_вес = round(nn_dohodi_исполнение / (form_data['доходы_исполненный'] / 1000) * 100, 1)
            
            replacements['<NNDohodi>'] = self._format_number(nn_dohodi_исполнение)
            replacements['<NNPD>'] = self._format_number(nn_dohodi_план, decimals=1)
            replacements['<ObshiiPr>'] = self._format_number(nn_dohodi_удельный_вес, decimals=1)
        
        # Метки для безвозмездных поступлений
        if sbpb_item:
            sbpb_исполнение = sbpb_item.get('исполненный', {}).get(budget_col, 0) / 1000
            sbpb_план = 0
            sbpb_убн = sbpb_item.get('утвержденный', {}).get(budget_col, 0) / 1000
            if sbpb_убн != 0:
                sbpb_план = round(sbpb_исполнение / sbpb_убн * 100, 1)
            sbpb_удельный_вес = 0
            if form_data['доходы_исполненный'] != 0:
                sbpb_удельный_вес = round(sbpb_исполнение / (form_data['доходы_исполненный'] / 1000) * 100, 1)
            
            replacements['<SBpB>'] = self._format_number(sbpb_исполнение)
            replacements['<PrSBpB>'] = self._format_number(sbpb_план, decimals=1)
            replacements['<OPrSBpB>'] = self._format_number(sbpb_удельный_вес, decimals=1)
        
        # Инициализируем переменные для использования в метках
        nn_dohodi_исполнение = 0
        nn_dohodi_план = 0
        nn_dohodi_удельный_вес = 0
        sbpb_исполнение = 0
        
        # Метки для максимальных удельных весов (из Таблицы 2 и 3)
        if max_доход:
            replacements['<NDUr2>'] = max_доход.get('наименование', '')
            replacements['<SDUr2>'] = self._format_number(max_доход.get('исполнение', 0))
            # Процент от налоговых и неналоговых доходов (NNDohodi)
            nn_dohodi_total = nn_dohodi_исполнение if nn_dohodi_item else form_data['доходы_исполненный'] / 1000
            pr_sd_ur2 = 0
            if nn_dohodi_total != 0:
                pr_sd_ur2 = round(max_доход.get('исполнение', 0) / nn_dohodi_total * 100, 1)
            replacements['<PRSDUr2>'] = self._format_number(pr_sd_ur2, decimals=1)
        
        if max_расход:
            # Первая буква в нижнем регистре
            наименование = max_расход.get('наименование', '')
            if наименование:
                наименование = наименование[0].lower() + наименование[1:] if len(наименование) > 1 else наименование.lower()
            replacements['<NDUr2B>'] = наименование
            replacements['<SDUr2B>'] = self._format_number(max_расход.get('исполнение', 0))
            # Процент от общих расходов (SBpB - безвозмездные поступления)
            sbpb_total = sbpb_исполнение if sbpb_item else form_data['расходы_исполненный'] / 1000
            pr_sd_ur2b = 0
            if sbpb_total != 0:
                pr_sd_ur2b = round(max_расход.get('исполнение', 0) / sbpb_total * 100, 1)
            replacements['<OPSDUr2B>'] = self._format_number(pr_sd_ur2b, decimals=1)
        
        # Анализ расходов: находим расходы ниже и выше общего процента исполнения
        vsego_ras_nige = 0
        s_rashod_lit = 0
        gde_nige_rasxod = ""
        gde_vishe_rasxod = ""
        rashod_big = None
        max_udelny_ves_расходы_для_анализа = 0
        
        # Собираем данные для анализа расходов (только уровень 1)
        расходы_для_анализа = []
        for item in расходы_data:
            уровень = item.get('уровень', 0)
            if уровень == 1:
                наименование = item.get('наименование_показателя', '')
                if 'расходы, всего' not in наименование.lower() and 'результат' not in наименование.lower():
                    убн = item.get('утвержденный', {}).get(budget_col, 0) / 1000
                    исполнение = item.get('исполненный', {}).get(budget_col, 0) / 1000
                    план = round((исполнение / убн * 100) if убн != 0 else 0, 1)
                    удельный_вес = round((исполнение / (form_data['расходы_исполненный'] / 1000) * 100) if form_data['расходы_исполненный'] != 0 else 0, 1)
                    
                    расходы_для_анализа.append({
                        'наименование': наименование,
                        'план': план,
                        'удельный_вес': удельный_вес,
                        'исполнение': исполнение,
                        'код': item.get('код_классификации', '').replace(' ', '')
                    })
        
        # Находим максимальный удельный вес среди расходов уровня 1
        if расходы_для_анализа:
            max_udelny_ves_расходы_для_анализа = max(item['удельный_вес'] for item in расходы_для_анализа)
            
            for item in расходы_для_анализа:
                if item['план'] < form_data['процент_расходов']:
                    vsego_ras_nige += 1
                    gde_nige_rasxod += f"«{item['наименование']}» ({item['план']}%), "
                else:
                    s_rashod_lit += 1
                    gde_vishe_rasxod += f"«{item['наименование']}» ({item['план']}%), "
                
                if abs(item['удельный_вес'] - max_udelny_ves_расходы_для_анализа) < 0.01:
                    код_расхода = item['код']
                    код_подраздела = код_расхода[7:11] if len(код_расхода) >= 11 else ''
                    rashod_big = {
                        'наименование': f"{код_подраздела} {item['наименование']}",
                        'исполнение': item['исполнение'],
                        'удельный_вес': item['удельный_вес']
                    }
        
        # Очищаем строки от лишних запятых
        gde_nige_rasxod = gde_nige_rasxod.rstrip(', ').replace('%), .', '%).')
        gde_vishe_rasxod = gde_vishe_rasxod.rstrip(', ').replace('%), .', '%).')
        
        replacements['<VsegoRasNige>'] = str(vsego_ras_nige)
        replacements['<SRashodLit>'] = str(s_rashod_lit)
        replacements['<GdeNigeRasxod>'] = gde_nige_rasxod
        replacements['<GdeVisheRasxod>'] = gde_vishe_rasxod
        
        if rashod_big:
            replacements['<RashodBig>'] = rashod_big['наименование']
            replacements['<SRashodBig>'] = self._format_number(rashod_big['исполнение'])
            replacements['<PrVsexRas>'] = self._format_number(rashod_big['удельный_вес'], decimals=1)
        
        # Формируем текст Vibor1 на основе сравнения процентов исполнения
        vibor1 = ""
        if nn_dohodi_item and len(расходы_для_анализа) > 0:
            первый_план = расходы_для_анализа[0]['план'] if расходы_для_анализа else 0
            последний_план = form_data['процент_доходов']
            if abs(первый_план - последний_план) < 10:
                vibor1 = (f"При общем выполнении доходной части бюджета по налоговым и неналоговым доходам "
                         f"на {self._format_number(nn_dohodi_удельный_вес, decimals=1)}% от утвержденных бюджетных назначений, "
                         f"по отдельным видам доходов уровень исполнения за {period_name} {protocol_date.year} года "
                         f"значительно превышает ¾ годовых плановых показателей.")
            else:
                vibor1 = (f"Выполнение доходной части бюджета по налоговым и неналоговым доходам за "
                         f"{period_name} {protocol_date.year} года значительно превышает ¾ годовых плановых показателей.")
        replacements['<Vibor1>'] = vibor1
        
        # Метка PsT2 (процент исполнения для кода "10000000000000000")
        if nn_dohodi_item:
            replacements['<PsT2>'] = self._format_number(nn_dohodi_план, decimals=1)
        
        # Заменяем метки в тексте документа
        for paragraph in doc.paragraphs:
            for key, value in replacements.items():
                if key in paragraph.text:
                    paragraph.text = paragraph.text.replace(key, str(value))
        
        # Заменяем метки в таблицах
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in replacements.items():
                        if key in cell.text:
                            cell.text = cell.text.replace(key, str(value))
    
    def _insert_table2(self, doc: docx.Document, form_data: Dict[str, Any], protocol_date: datetime):
        """Вставка Таблицы 2 (доходы) - логика из 1С"""
        # Ищем маркер "Таблица 2" в документе
        table2_marker = None
        for i, paragraph in enumerate(doc.paragraphs):
            if 'Таблица 2' in paragraph.text:
                table2_marker = i
                break
        
        if table2_marker is None:
            logger.warning("Маркер 'Таблица 2' не найден в документе")
            return
        
        # Загружаем справочник доходов для определения уровня
        income_reference_df = self.db_manager.load_income_reference_df()
        
        # Формируем данные для таблицы
        доходы_data = form_data['доходы_data']
        table_data = []
        budget_col = 'бюджет субъекта Российской Федерации'
        доходы_исполненный_всего = form_data['доходы_исполненный']
        
        # Фильтруем доходы по уровням (уровень = 1, 2, 3) согласно логике 1С
        for item in доходы_data:
            код = item.get('код_классификации', '').replace(' ', '')
            if not код or len(код) < 20:
                continue
            
            # Используем уровень из данных формы
            уровень = item.get('уровень', 0)
            
            # Если уровня нет в данных, пытаемся получить из справочника
            if уровень == 0 and not income_reference_df.empty:
                уровень = self._get_level_from_reference(код, income_reference_df)
            
            # Включаем только строки с уровнем = 1, 2 или 3
            if уровень in [1, 2, 3]:
                убн = item.get('утвержденный', {}).get(budget_col, 0) / 1000
                исполнение = item.get('исполненный', {}).get(budget_col, 0) / 1000
                
                # Вычисляем процент исполнения
                план = 0
                if убн != 0:
                    план = round(исполнение / убн * 100, 1)
                
                # Вычисляем удельный вес
                удельный_вес = 0
                if доходы_исполненный_всего != 0:
                    удельный_вес = round(исполнение / (доходы_исполненный_всего / 1000) * 100, 1)
                
                table_data.append({
                    'наименование': item.get('наименование_показателя', ''),
                    'убн': убн,
                    'исполнение': исполнение,
                    'план': план,
                    'удельный_вес': удельный_вес,
                    'код': код,
                    'уровень': уровень
                })
        
        # Добавляем итоговую строку
        table_data.append({
            'наименование': 'Доходы, всего',
            'убн': form_data['доходы_утвержденный'] / 1000,
            'исполнение': form_data['доходы_исполненный'] / 1000,
            'план': form_data['процент_доходов'],
            'удельный_вес': 100.0,
            'код': '',
            'уровень': 1
        })
        
        # Находим позицию для вставки таблицы (после маркера)
        paragraph = doc.paragraphs[table2_marker]
        table = doc.add_table(rows=len(table_data) + 1, cols=4)
        # Используем стандартный стиль таблицы или создаем без стиля
        style_applied = False
        try:
            table.style = 'Light Grid Accent 1'
            style_applied = True
        except:
            # Если стиль не найден, используем базовый стиль или форматируем вручную
            try:
                table.style = 'Table Grid'
                style_applied = True
            except:
                # Если и этот стиль не найден, применяем форматирование вручную
                self._apply_table_borders(table)
        
        # Настройка стилей заголовков
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Наименование показателя'
        header_cells[1].text = f'Утвержденные бюджетные назначения на {protocol_date.year} год, тыс. рублей'
        header_cells[2].text = f'Исполнение на {protocol_date.strftime("%d.%m.%Y")}, тыс. рублей'
        header_cells[3].text = 'Уровень исполнения, %'
        
        # Применяем стили к заголовкам
        for cell in header_cells:
            if cell.paragraphs and cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Данные
        for i, row_data in enumerate(table_data, start=1):
            cells = table.rows[i].cells
            cells[0].text = row_data['наименование']
            cells[1].text = self._format_number(row_data['убн'])
            cells[2].text = self._format_number(row_data['исполнение'])
            cells[3].text = self._format_number(row_data['план'], decimals=1)
            
            # Жирный шрифт для строк с уровнем = 1 или 2
            if row_data['уровень'] in [1, 2]:
                for cell in cells:
                    if cell.paragraphs and cell.paragraphs[0].runs:
                        cell.paragraphs[0].runs[0].font.bold = True
    
    def _insert_table3(self, doc: docx.Document, form_data: Dict[str, Any], protocol_date: datetime):
        """Вставка Таблицы 3 (расходы) - логика из 1С"""
        # Ищем маркер "Таблица 3" в документе
        table3_marker = None
        for i, paragraph in enumerate(doc.paragraphs):
            if 'Таблица 3' in paragraph.text:
                table3_marker = i
                break
        
        if table3_marker is None:
            logger.warning("Маркер 'Таблица 3' не найден в документе")
            return
        
        # Формируем данные для таблицы
        расходы_data = form_data['расходы_data']
        table_data = []
        budget_col = 'бюджет субъекта Российской Федерации'
        расходы_исполненный_всего = form_data['расходы_исполненный']
        
        # Фильтруем расходы по уровням (уровень = 1, 2) согласно логике 1С
        for item in расходы_data:
            код = item.get('код_классификации', '').replace(' ', '')
            if not код or len(код) != 20:
                continue
            
            # Используем уровень из данных формы
            уровень = item.get('уровень', 0)
            
            # Если уровня нет, определяем по структуре кода
            if уровень == 0:
                if код[3:5] != '00':
                    уровень = 1
                if уровень == 1 and код[5:7] != '00':
                    уровень = 2
            
            # Пропускаем строки с "x" в коде (служебные)
            if 'x' in item.get('наименование_показателя', '').lower():
                continue
            
            # Включаем только строки с уровнем = 1 или 2
            if уровень in [1, 2]:
                убн = item.get('утвержденный', {}).get(budget_col, 0) / 1000
                исполнение = item.get('исполненный', {}).get(budget_col, 0) / 1000
                
                # Вычисляем процент исполнения
                план = 0
                if убн != 0:
                    план = round(исполнение / убн * 100, 1)
                
                # Вычисляем удельный вес
                удельный_вес = 0
                if расходы_исполненный_всего != 0:
                    удельный_вес = round(исполнение / (расходы_исполненный_всего / 1000) * 100, 1)
                
                # Формируем код раздела/подраздела (первые 4 символа после первых 3)
                код_рп_форматированный = код[3:7] if len(код) >= 7 else ''
                
                table_data.append({
                    'наименование': item.get('наименование_показателя', ''),
                    'код_рп': код_рп_форматированный,
                    'убн': убн,
                    'исполнение': исполнение,
                    'план': план,
                    'удельный_вес': удельный_вес,
                    'код': код,
                    'уровень': уровень
                })
        
        # Добавляем итоговую строку
        table_data.append({
            'наименование': 'РАСХОДЫ, ВСЕГО',
            'код_рп': '',
            'убн': form_data['расходы_утвержденный'] / 1000,
            'исполнение': form_data['расходы_исполненный'] / 1000,
            'план': form_data['процент_расходов'],
            'удельный_вес': 100.0,
            'код': '',
            'уровень': 1
        })
        
        # Создаем таблицу после маркера
        paragraph = doc.paragraphs[table3_marker]
        table = doc.add_table(rows=len(table_data) + 1, cols=5)
        # Используем стандартный стиль таблицы или создаем без стиля
        style_applied = False
        try:
            table.style = 'Light Grid Accent 1'
            style_applied = True
        except:
            # Если стиль не найден, используем базовый стиль или форматируем вручную
            try:
                table.style = 'Table Grid'
                style_applied = True
            except:
                # Если и этот стиль не найден, применяем форматирование вручную
                self._apply_table_borders(table)
        
        # Настройка стилей заголовков
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Наименование показателя'
        header_cells[1].text = 'Раздел, подраздел'
        header_cells[2].text = f'Утвержденные бюджетные назначения на {protocol_date.year} год, тыс. рублей'
        header_cells[3].text = f'Исполнение на {protocol_date.strftime("%d.%m.%Y")}, тыс. рублей'
        header_cells[4].text = 'Процент исполнения, %'
        
        # Применяем стили к заголовкам
        for cell in header_cells:
            if cell.paragraphs and cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Данные
        for i, row_data in enumerate(table_data, start=1):
            cells = table.rows[i].cells
            cells[0].text = row_data['наименование']
            cells[1].text = row_data['код_рп']
            cells[2].text = self._format_number(row_data['убн'])
            cells[3].text = self._format_number(row_data['исполнение'])
            cells[4].text = self._format_number(row_data['план'], decimals=1)
            
            # Жирный шрифт для строк с уровнем = 1
            if row_data['уровень'] == 1:
                for cell in cells:
                    if cell.paragraphs and cell.paragraphs[0].runs:
                        cell.paragraphs[0].runs[0].font.bold = True
    
    def _generate_letter_admin(
        self,
        template_path: Path,
        project: Project,
        municipality: Optional[Any],
        protocol_date: datetime,
        protocol_number: str,
        output_dir: Optional[str]
    ) -> Optional[str]:
        """Формирование письма администрации"""
        try:
            doc = docx.Document(str(template_path))
            
            replacements = {
                '<Dolzhnost>': 'Главе муниципального образования',
                '<Okrug1>': municipality.name if municipality else project.name,
                '<FIO>': self._format_fio(
                    getattr(municipality, 'фамилия_администрация', ''),
                    getattr(municipality, 'имя_администрация', ''),
                    getattr(municipality, 'отчество_администрация', '')
                ),
                '<Adr>': getattr(municipality, 'адрес_администрация', ''),
                '<Email>': getattr(municipality, 'администрация_почта', ''),
                '<DatePr>': protocol_date.strftime('%d.%m.%Y'),
                '<NomerPr>': protocol_number,
            }
            
            self._replace_in_document(doc, replacements)
            
            if not output_dir:
                output_dir = Path("data") / "projects" / str(project.id) / "documents"
            output_dir = Path(output_dir)
            output_dir.mkdir(parents=True, exist_ok=True)
            
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = output_dir / f"письмо_администрации_{timestamp}.docx"
            doc.save(str(output_path))
            
            return str(output_path)
            
        except Exception as e:
            logger.error(f"Ошибка формирования письма администрации: {e}", exc_info=True)
            return None
    
    def _generate_letter_council(
        self,
        template_path: Path,
        project: Project,
        municipality: Optional[Any],
        protocol_date: datetime,
        protocol_number: str,
        output_dir: Optional[str]
    ) -> Optional[str]:
        """Формирование письма совета"""
        try:
            doc = docx.Document(str(template_path))
            
            okrug_rp = getattr(municipality, 'родительный_падеж', municipality.name if municipality else project.name)
            okrug_council = okrug_rp.replace('округа', 'совета')
            
            replacements = {
                '<Dolzhnost>': 'Председателю',
                '<Okrug1>': okrug_council,
                '<FIO>': self._format_fio(
                    getattr(municipality, 'фамилия_совет', ''),
                    getattr(municipality, 'имя_совет', ''),
                    getattr(municipality, 'отчество_совет', '')
                ),
                '<Adr>': getattr(municipality, 'адрес_совет', ''),
                '<Email>': getattr(municipality, 'совет_почта', ''),
                '<DatePr>': protocol_date.strftime('%d.%m.%Y'),
                '<NomerPr>': protocol_number,
            }
            
            self._replace_in_document(doc, replacements)
            
            if not output_dir:
                output_dir = Path("data") / "projects" / str(project.id) / "documents"
            output_dir = Path(output_dir)
            output_dir.mkdir(parents=True, exist_ok=True)
            
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = output_dir / f"письмо_совета_{timestamp}.docx"
            doc.save(str(output_path))
            
            return str(output_path)
            
        except Exception as e:
            logger.error(f"Ошибка формирования письма совета: {e}", exc_info=True)
            return None
    
    def _replace_in_document(self, doc: docx.Document, replacements: Dict[str, str]):
        """Замена меток в документе"""
        for paragraph in doc.paragraphs:
            for key, value in replacements.items():
                if key in paragraph.text:
                    paragraph.text = paragraph.text.replace(key, str(value))
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in replacements.items():
                        if key in cell.text:
                            cell.text = cell.text.replace(key, str(value))
    
    def _format_number(self, value: float, decimals: int = 1) -> str:
        """Форматирование числа с разделителями"""
        if value == 0:
            return "0,0"
        formatted = f"{value:,.{decimals}f}".replace(',', ' ').replace('.', ',')
        return formatted
    
    def _format_fio(self, surname: str, name: str, patronymic: str) -> str:
        """Форматирование ФИО в формате: Фамилия И.О."""
        if not surname:
            return ""
        initials = ""
        if name:
            initials += name[0] + "."
        if patronymic:
            initials += patronymic[0] + "."
        return f"{surname} {initials}".strip()
    
    def _get_level_from_reference(self, code: str, reference_df) -> int:
        """Получение уровня из справочника доходов"""
        try:
            import pandas as pd
            if reference_df is None or reference_df.empty:
                return 0
            
            # Ищем код в справочнике
            match = reference_df[reference_df['код_классификации_ДБ'] == code]
            if not match.empty:
                level = match.iloc[0]['уровень_кода']
                return int(level) if pd.notna(level) else 0
        except Exception as e:
            logger.warning(f"Ошибка получения уровня из справочника: {e}")
        
        return 0
    
    def _find_max_udelny_ves(self, data: List[Dict[str, Any]], total_executed: float) -> float:
        """Нахождение максимального удельного веса в данных"""
        max_udelny_ves = 0.0
        budget_col = 'бюджет субъекта Российской Федерации'
        
        for item in data:
            код = item.get('код_классификации', '').replace(' ', '')
            уровень = item.get('уровень', 0)
            
            # Для доходов: только уровень 3 и не начинается с "2" в 5-й позиции
            # Для расходов: только уровень 1 или 2
            if len(код) >= 20:
                if len(data) > 0 and 'доходы' in str(data[0].get('раздел', '')).lower():
                    if уровень == 3 and (len(код) < 5 or код[4] != '2'):
                        исполнение = item.get('исполненный', {}).get(budget_col, 0) / 1000
                        if total_executed != 0:
                            удельный_вес = round(исполнение / (total_executed / 1000) * 100, 1)
                            if удельный_вес > max_udelny_ves:
                                max_udelny_ves = удельный_вес
                elif len(data) > 0 and 'расходы' in str(data[0].get('раздел', '')).lower():
                    if уровень in [1, 2]:
                        исполнение = item.get('исполненный', {}).get(budget_col, 0) / 1000
                        if total_executed != 0:
                            удельный_вес = round(исполнение / (total_executed / 1000) * 100, 1)
                            if удельный_вес > max_udelny_ves:
                                max_udelny_ves = удельный_вес
        
        return max_udelny_ves
    
    def _find_item_by_udelny_ves(self, data: List[Dict[str, Any]], target_udelny_ves: float, total_executed: float) -> Optional[Dict[str, Any]]:
        """Нахождение элемента с заданным удельным весом"""
        budget_col = 'бюджет субъекта Российской Федерации'
        
        for item in data:
            код = item.get('код_классификации', '').replace(' ', '')
            уровень = item.get('уровень', 0)
            
            if len(код) >= 20:
                исполнение = item.get('исполненный', {}).get(budget_col, 0) / 1000
                if total_executed != 0:
                    удельный_вес = round(исполнение / (total_executed / 1000) * 100, 1)
                    if abs(удельный_вес - target_udelny_ves) < 0.01:  # Сравнение с небольшой погрешностью
                        return {
                            'наименование': item.get('наименование_показателя', ''),
                    'исполнение': исполнение,
                    'удельный_вес': удельный_вес
                }
        
        return None
    
    def _apply_table_borders(self, table):
        """Применение границ таблицы вручную, если стиль не применился"""
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        
        tbl = table._tbl
        tblBorders = OxmlElement('w:tblBorders')
        
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            tblBorders.append(border)
        
        tbl.tblPr.append(tblBorders)

